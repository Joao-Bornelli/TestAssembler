from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

import json

def generate_docx(fillData, questionsData):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    sections = doc.sections[0]
    sections.top_margin = Cm(1.0)
    sections.bottom_margin = Cm(1.0)
    sections.left_margin = Cm(1.0)
    sections.right_margin = Cm(1.0)

    #Cabeçalho da prova
    table = doc.add_table(rows=7,cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = "Table Grid"
    table.autofit = False
    
    table.columns[0].width = Inches(2.0)
    table.columns[1].width = Inches(4.0)
    table.columns[2].width = Inches(1.5)

    header = [table.cell(0,1),table.cell(0,2)]
    header[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # ====
    #	Coluna da esquerda, com o logo
    # ====

    left_column_cell = table.cell(0,0).merge(table.cell(6,0))
    left_column_cell.text = ''
    left_column_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    left_column_image = left_column_cell.add_paragraph()
    left_column_image.alignment = WD_ALIGN_PARAGRAPH.CENTER

    left_column_run_img = left_column_image.add_run()
    left_column_run_img.add_picture("include/LogoSenai.png",width=Inches(1.5))

    left_column_text = left_column_cell.add_paragraph()
    left_column_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    left_column_text.space_before = Pt(10)
    left_column_text.space_after = Pt(10)

    left_column_run_text = left_column_text.add_run("Serviço Nacional de Aprendizagem Industrial")
    left_column_run_text.bold = True
    left_column_run_text = left_column_text.add_run("\nSanta Catarina")

    # ====
    #   Coluna da direita, com o desempenho
    # ====

    rightColumn = table.cell(0,2).merge(table.cell(6,2))
    rightColumn.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    header[1] = header[1].paragraphs[0]
    header[1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header[1].vertical_alignment = WD_ALIGN_VERTICAL.TOP
    header[1].text = ''
    header_run = header[1].add_run("Desempenho")
    header_run.bold = True

    # ====
    #   Coluna central com as infos
    # ====
    header[0] = header[0].paragraphs[0]
    header[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    header[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    header[0].text = ''
    header_run2 = header[0].add_run("Avaliação Objetiva")
    header_run2.bold = True


    data = {'Date'		 :table.cell(1,1),
            'Instructor' :table.cell(2,1),
            'Course'	 :table.cell(3,1),
            'Course Unit':table.cell(4,1),
            'Class'		 :table.cell(5,1),
            'Student'	 :table.cell(6,1),}
    labels = ['Data: ','Docente: ','Curso: ','Unidade Curricular: ','Turma: ','Estudante: ']
    values= ['___/___/_____',           #Data
            fillData['Instructor'],     #Docente   
            fillData['Course'],         #Curso
            fillData['Course_Unit'],    #Unidade Curricular
            fillData['Class'],          #Turma
            '']                         #Estudante

    print('AQUI')

    for i, key in enumerate(data.keys()):
        cell = data[key]                            # mantém a célula original
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ''                              # remove parágrafo automático
        p = cell.paragraphs[0]                      # usa o parágrafo existente
        r = p.add_run(labels[i] + values[i])        # adiciona texto
        r.bold = True


    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(12)
    
    try:
        with open('include/data.json', 'r')as file:
            jsonData = json.load(file)
    except Exception as e:
        print("Erro ao carregar o arquivo JSON:", e)

    for i, question in enumerate(jsonData):
        
        question_table = doc.add_table(rows=3,cols=1)
        question_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        question_table.style = "Table Grid"
        question_table.autofit = True
        question_table.columns[0].width = Inches(7.5)        
        
        item_cell = question_table.cell(0,0)
        item_cell.text = ""                    
        item_row = item_cell.paragraphs[0]
        
        item_row.add_run(f"ITEM {i+1}")
        set_cell_background(item_cell,'00418a')
        
        hability_cell = question_table.cell(1,0)
        hability_cell.text = ""
        hability_row = hability_cell.paragraphs[0]
        hability_row.add_run("Habilidade: ").bold = True
        hability_row.add_run(question['Informações Essenciais']['Capacidade avaliada']).bold = False
        
        context_command_cell = question_table.cell(2,0)
        context_command_cell.text = ""
        context_paragraph = context_command_cell.paragraphs[0]
        context_paragraph.add_run("Contexto: ").bold = True
        context_paragraph.add_run(question['Contexto']).bold = False
        
        command_paragraph = context_command_cell.add_paragraph()
        command_paragraph.add_run("\nComando: ").bold = True
        command_paragraph.add_run(question['Comando']).bold = False
        
        alternatives_paragraph = context_command_cell.add_paragraph()
        alternatives_paragraph.add_run("\nAlternativas:").bold = True
        for j, alternative in enumerate(question['Alternativas']):
            alternatives_paragraph.add_run(f'\n{chr(65+j)}. ').bold = True
            alternatives_paragraph.add_run(alternative['texto']).bold = False        
        
        for row in question_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.keep_together = True
                    p.paragraph_format.keep_with_next = True
        
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(12)
        
    doc.save("arquivo.docx")
    
def set_cell_background(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

data = {'Instructor': 'João Bornelli',
        'Course': 'Técnico em Informática',
        'Course_Unit': 'Desenvolvimento de Sistemas',
        'Class': '1A'
        }

generate_docx(data,[])
